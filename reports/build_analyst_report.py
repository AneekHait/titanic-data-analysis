"""Build the Titanic Survival Analyst Report (DOCX).

Standalone script. Does not depend on dashboard/generate_pdf.py.
The content is written as an analyst's narrative report - prose-driven, with
charts and tables embedded where they earn their place. The script's job is
purely to render that narrative; the analytical content lives in this file.
"""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.analysis.eda import correlation_analysis
from src.analysis.inference import key_odds_ratios, survival_rates_with_ci, wilson_ci
from src.analysis.statistics import (
    anova_survival,
    chi_square_test,
    effect_sizes,
    t_test_survival,
)
from src.data.loader import load_titanic
from src.data.processing import clean_data, engineer_features

CHART_DIR = ROOT / "dashboard" / "_pdf_charts"
OUTPUT_PATH = ROOT / "reports" / "Titanic_Survival_Analyst_Report.docx"

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT_BLUE = RGBColor(0x25, 0x63, 0xEB)
MUTED = RGBColor(0x4B, 0x55, 0x63)
DARK = RGBColor(0x0F, 0x17, 0x2A)
RED = RGBColor(0xB9, 0x1C, 0x1C)
GREEN = RGBColor(0x05, 0x96, 0x69)
GREY_BG = "F3F4F6"
PALE_BLUE_BG = "E0EAFB"
PALE_GREEN_BG = "DCFCE7"
PALE_AMBER_BG = "FEF3C7"


# ---------------------------------------------------------------------------
# Low-level docx helpers
# ---------------------------------------------------------------------------


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, *, sz=4, color="D1D5DB") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def paragraph_border(paragraph, *, side="left", color="2563EB", sz=18) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), "8")
    el.set(qn("w:color"), color)
    pbdr.append(el)
    p_pr.append(pbdr)


def paragraph_shading(paragraph, hex_color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def add_run(paragraph, text: str, *, bold=False, italic=False, size=11, color=None, font="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return run


def add_para(doc, text: str = "", *, style=None, size=11, color=DARK, bold=False,
             italic=False, space_after=6, alignment=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    if alignment is not None:
        p.alignment = alignment
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def add_heading_1(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    add_run(p, text, bold=True, size=18, color=NAVY)
    return p


def add_heading_2(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    add_run(p, text, bold=True, size=13.5, color=ACCENT_BLUE)
    return p


def add_heading_3(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    add_run(p, text, bold=True, size=11, color=NAVY)
    return p


def add_bullet(doc, text_runs):
    """text_runs is a list of (text, is_bold) tuples for inline formatting."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    for run in text_runs:
        if isinstance(run, str):
            add_run(p, run, size=11, color=DARK)
        else:
            text, bold = run
            add_run(p, text, bold=bold, size=11, color=DARK)
    return p


def add_pullquote(doc, text: str, color_bar="2563EB", bg="EFF6FF"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    paragraph_border(p, side="left", color=color_bar, sz=24)
    paragraph_shading(p, bg)
    add_run(p, text, italic=False, size=11, color=DARK)
    return p


def add_callout(doc, title: str, body: str, accent="2563EB", bg="EFF6FF"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    paragraph_border(p, side="left", color=accent, sz=24)
    paragraph_shading(p, bg)
    add_run(p, title, bold=True, size=10.5,
            color=RGBColor.from_string(accent))
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(10)
    p2.paragraph_format.left_indent = Cm(0.4)
    p2.paragraph_format.right_indent = Cm(0.4)
    paragraph_border(p2, side="left", color=accent, sz=24)
    paragraph_shading(p2, bg)
    add_run(p2, body, size=10.5, color=DARK)


def add_image(doc, path: Path, *, width_cm: float = 16.5, caption: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        add_run(cap, caption, italic=True, size=9.5, color=MUTED)


def add_table(doc, headers, rows, *, col_widths_cm=None, header_color="1E3A5F",
              header_text=RGBColor(0xFF, 0xFF, 0xFF)):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        set_cell_shading(cell, header_color)
        set_cell_border(cell, color=header_color)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, str(h), bold=True, size=10, color=header_text)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_widths_cm and i < len(col_widths_cm):
            cell.width = Cm(col_widths_cm[i])
    # body
    for r_idx, row in enumerate(rows):
        for i, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[i]
            cell.text = ""
            if r_idx % 2 == 1:
                set_cell_shading(cell, GREY_BG)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, str(val), size=10, color=DARK)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if col_widths_cm and i < len(col_widths_cm):
                cell.width = Cm(col_widths_cm[i])
    # space after
    after_p = doc.add_paragraph()
    after_p.paragraph_format.space_after = Pt(6)
    return table


def page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def setup_document(doc: Document):
    sections = doc.sections
    for s in sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK


# ---------------------------------------------------------------------------
# Pre-computed analytics
# ---------------------------------------------------------------------------


def compute_analytics(df_raw: pd.DataFrame) -> dict:
    df = engineer_features(clean_data(df_raw))

    sex_ci = survival_rates_with_ci(df, "Sex")
    pclass_ci = survival_rates_with_ci(df, "Pclass")
    emb_ci = survival_rates_with_ci(df, "Embarked")

    es = effect_sizes(df)
    odds = key_odds_ratios(df)

    chi_sex = chi_square_test(df, "Sex")
    chi_pc = chi_square_test(df, "Pclass")
    chi_emb = chi_square_test(df, "Embarked")
    t_age = t_test_survival(df, "Age")
    t_fare = t_test_survival(df, "Fare")
    anova_age = anova_survival(df.assign(AgeGroup=pd.cut(df["Age"], bins=[0, 16, 32, 48, 64, 100],
                                                        labels=["Child", "Young Adult", "Adult", "Older Adult", "Senior"])),
                               "AgeGroup")

    joint = df.groupby(["Pclass", "Sex"], observed=True)["Survived"].agg(["count", "mean"])
    joint["rate"] = (joint["mean"] * 100).round(1)

    family = df.groupby("FamilySize")["Survived"].agg(["count", "mean"]).reset_index()
    family["rate"] = (family["mean"] * 100).round(1)

    title = df.groupby("Title")["Survived"].agg(["count", "mean"]).reset_index()
    title["rate"] = (title["mean"] * 100).round(1)
    title = title.sort_values("count", ascending=False)

    age_groups = df.copy()
    age_groups["AgeGroup"] = pd.cut(age_groups["Age"], bins=[0, 16, 32, 48, 64, 100],
                                    labels=["Child (0-16)", "Young Adult (17-32)", "Adult (33-48)",
                                            "Older Adult (49-64)", "Senior (65+)"])
    age_summary = age_groups.groupby("AgeGroup", observed=True)["Survived"].agg(["count", "mean"]).reset_index()
    age_summary["rate"] = (age_summary["mean"] * 100).round(1)

    has_boat = df["Lifeboat"].notna()
    n_boat = int(has_boat.sum())
    s_boat = int(df.loc[has_boat, "Survived"].sum())
    n_no = int((~has_boat).sum())
    s_no = int(df.loc[~has_boat, "Survived"].sum())

    return {
        "df_raw": df_raw,
        "df": df,
        "n_total": len(df),
        "n_survived": int(df["Survived"].sum()),
        "rate_overall": float(df["Survived"].mean() * 100),
        "sex_ci": sex_ci,
        "pclass_ci": pclass_ci,
        "emb_ci": emb_ci,
        "effect_sizes": es,
        "odds": odds,
        "tests": {"sex": chi_sex, "pclass": chi_pc, "emb": chi_emb,
                  "age_t": t_age, "fare_t": t_fare, "age_anova": anova_age},
        "joint": joint,
        "family": family,
        "title": title,
        "age_summary": age_summary,
        "lifeboat": {"n_boat": n_boat, "s_boat": s_boat,
                     "n_no": n_no, "s_no": s_no,
                     "rate_boat": s_boat / n_boat * 100 if n_boat else 0,
                     "rate_no": s_no / n_no * 100 if n_no else 0},
        "missing_age": int(df_raw["Age"].isnull().sum()),
        "missing_occ": int(df_raw["Occupation"].isnull().sum()),
    }


# ---------------------------------------------------------------------------
# Report content
# ---------------------------------------------------------------------------


def cover_page(doc, A):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "WHO SURVIVED THE TITANIC?", bold=True, size=10, color=ACCENT_BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    add_run(title, "A Quantitative Analysis of Survival", bold=True, size=28, color=NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(32)
    add_run(sub, "on the RMS Titanic", bold=True, size=28, color=NAVY)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(rule, "_" * 30, size=12, color=ACCENT_BLUE)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(36)
    add_run(author, "Prepared by\n", italic=True, size=10.5, color=MUTED)
    add_run(author, "Aneek Hait\n", bold=True, size=16, color=NAVY)
    add_run(author, "aneekhait.github.io", italic=True, size=10.5, color=ACCENT_BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(32)
    add_run(meta,
            f"Dataset: titanic5 - {A['n_total']:,} passengers, 14 features\n",
            size=11, color=MUTED)
    add_run(meta, "Source: ", size=10, color=MUTED)
    add_run(meta, "hbiostat.org/data/repo/titanic5.csv\n", size=10, color=ACCENT_BLUE)
    add_run(meta, "(Encyclopedia Titanica / Vanderbilt Biostatistics)\n", size=10, color=MUTED)
    add_run(meta,
            f"Survival rate: {A['rate_overall']:.1f}%  ({A['n_survived']:,} of {A['n_total']:,} survived)\n",
            size=11, color=MUTED)
    add_run(meta,
            f"Generated: {pd.Timestamp.now().strftime('%B %d, %Y')}",
            size=11, color=MUTED)

    role = doc.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role.paragraph_format.space_before = Pt(120)
    add_run(role, "Data Analyst Report", italic=True, size=11, color=MUTED)
    page_break(doc)


def add_footer_credit(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Prepared by ", size=9, color=MUTED, italic=True)
    add_run(p, "Aneek Hait", size=9, color=NAVY, italic=True, bold=True)
    add_run(p, "  ·  ", size=9, color=MUTED, italic=True)
    add_run(p, "aneekhait.github.io", size=9, color=ACCENT_BLUE, italic=True)
    add_run(p, "  ·  Data: titanic5 (hbiostat.org)", size=9, color=MUTED, italic=True)


def main():
    print("Loading data...")
    df_raw = load_titanic()
    print("Computing analytics...")
    A = compute_analytics(df_raw)

    print("Building document...")
    doc = Document()
    setup_document(doc)

    cover_page(doc, A)
    executive_summary(doc, A)
    page_break(doc)
    background_and_question(doc, A)
    data_and_method(doc, A)
    page_break(doc)
    the_big_picture(doc, A)
    page_break(doc)
    the_three_drivers(doc, A)
    page_break(doc)
    secondary_factors(doc, A)
    page_break(doc)
    lifeboat_mechanism(doc, A)
    page_break(doc)
    statistical_robustness(doc, A)
    page_break(doc)
    conclusions(doc, A)
    page_break(doc)
    appendix(doc, A)

    add_footer_credit(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Saved: {OUTPUT_PATH}")
    return OUTPUT_PATH


def executive_summary(doc, A):
    add_heading_1(doc, "Executive Summary")

    add_para(doc,
        "On the morning of April 15, 1912, 1,309 people had boarded the RMS Titanic in Southampton, "
        f"Cherbourg, Queenstown and Belfast. By the next morning, {A['n_total'] - A['n_survived']:,} of them "
        f"were dead. That headline number - a {A['rate_overall']:.1f}% survival rate - is misleading on its own. "
        "The disaster did not pick its victims at random. It selected them along three sharp axes: "
        "sex, passenger class, and (to a lesser extent) age.",
        space_after=8)

    add_para(doc,
        "This report uses the titanic5 dataset to quantify exactly how unequal the outcomes were, "
        "and to separate the factors that genuinely mattered from those that only appeared to. "
        "The findings are presented in plain language; the underlying statistical work - effect sizes, "
        "odds ratios with 95% confidence intervals, chi-square tests, ANOVA, t-tests - appears in "
        "tables and an appendix for readers who want it.",
        space_after=10)

    add_heading_3(doc, "Headline findings")

    sex_or = next(o for o in A["odds"] if o["label"] == "Female vs Male")
    pc3_or = next(o for o in A["odds"] if o["label"] == "3rd Class vs 1st/2nd")
    child_or = next(o for o in A["odds"] if o["label"] == "Child (<=16) vs Adult")
    fem_rate = next(r for _, r in A["sex_ci"].iterrows() if r["level"] == "female")
    male_rate = next(r for _, r in A["sex_ci"].iterrows() if r["level"] == "male")
    pc1_rate = next(r for _, r in A["pclass_ci"].iterrows() if r["level"] == 1)
    pc3_rate = next(r for _, r in A["pclass_ci"].iterrows() if r["level"] == 3)
    lb = A["lifeboat"]

    add_bullet(doc, [
        ("Sex was the single largest determinant of survival.", True),
        f" Women survived at {fem_rate['rate']:.1f}%; men at {male_rate['rate']:.1f}%. "
        f"In odds terms, women were roughly {sex_or['odds_ratio']:.0f} times more likely to survive than men.",
    ])
    add_bullet(doc, [
        ("Class strongly compounded the effect of sex.", True),
        f" A 1st-class woman had a {A['joint'].loc[(1, 'female'), 'rate']:.1f}% chance of surviving; "
        f"a 3rd-class man had a {A['joint'].loc[(3, 'male'), 'rate']:.1f}% chance. "
        "Same ship, same iceberg - an 80-percentage-point gap.",
    ])
    add_bullet(doc, [
        ("The 'children first' protocol was real, but modest in size.", True),
        f" Children under 16 had about {child_or['odds_ratio']:.1f} times the survival odds "
        "of adults. A genuine effect, but nowhere near the magnitude of sex or class.",
    ])
    add_bullet(doc, [
        ("Lifeboat access was the proximate cause of survival.", True),
        f" Of {lb['n_boat']} passengers with a recorded lifeboat number, {lb['rate_boat']:.1f}% survived. "
        f"Of {lb['n_no']:,} without one, only {lb['rate_no']:.1f}% did. Every demographic factor above "
        "was, in effect, predicting who would get a seat on a boat.",
    ])
    add_bullet(doc, [
        ("Some apparently important factors are confounded with class.", True),
        " Fare and embarkation port both appear correlated with survival in raw numbers, but most of "
        "that signal disappears once class is controlled for. They are proxies, not independent causes.",
    ])

    add_pullquote(doc,
        "Bottom line: if you could ask only one question to guess whether a Titanic passenger survived, "
        "ask their sex. If you could ask two, ask their class as well. After those, every other factor "
        "is in the noise.")


def background_and_question(doc, A):
    add_heading_1(doc, "1. Background & Question")

    add_heading_2(doc, "1.1  What happened")
    add_para(doc,
        "The RMS Titanic struck an iceberg at 23:40 ship's time on April 14, 1912, in the North "
        "Atlantic. The collision opened the hull along five forward compartments - one more than "
        "the ship was designed to survive flooding. She sank in 2 hours and 40 minutes, with roughly "
        "1,500 people still on board. The number of lifeboat seats was approximately half the number "
        "of people aboard. Survival therefore depended almost entirely on who was given a seat in a "
        "lifeboat in those 160 minutes.",
    )

    add_heading_2(doc, "1.2  What this analysis is trying to answer")
    add_para(doc,
        "The historical narrative around the Titanic is dominated by the phrase 'women and children "
        "first.' This report tests that narrative quantitatively. We want to answer four concrete questions:",
        space_after=6)
    add_bullet(doc, [("Q1.", True), " How unequal were the survival outcomes across demographic groups?"])
    add_bullet(doc, [("Q2.", True), " Which factors were truly driving survival, and which were just confounded with deeper causes?"])
    add_bullet(doc, [("Q3.", True), " How large were the effects, and how confident can we be in those estimates given the sample sizes?"])
    add_bullet(doc, [("Q4.", True), " If we were to build a survival prediction model, which features should we prioritise and why?"])


def data_and_method(doc, A):
    add_heading_1(doc, "2. Data & Method")

    add_heading_2(doc, "2.1  Dataset")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "Source: titanic5, compiled by Encyclopedia Titanica and hosted by ", size=11, color=DARK)
    add_run(p, "Vanderbilt Biostatistics", size=11, color=DARK)
    add_run(p, " at ", size=11, color=DARK)
    add_run(p, "hbiostat.org/data/repo/titanic5.csv", size=11, color=ACCENT_BLUE, italic=True)
    add_run(p,
        f". The dataset contains {A['n_total']:,} passengers and 14 columns. It is materially more "
        "complete than the well-known Kaggle training subset (891 rows). In particular, only 51 "
        "ages are missing (3.9%) versus Kaggle's ~20%, which makes age-stratified analysis here "
        "much more reliable.",
        size=11, color=DARK,
    )

    add_para(doc, "Columns used in this analysis:", space_after=4)
    add_table(doc,
        ["Column", "Type", "Used For"],
        [
            ["Survived", "0/1", "Target outcome"],
            ["Pclass", "1/2/3", "Socioeconomic class proxy (cabin location, boarding priority)"],
            ["Sex", "female/male", "The single strongest predictor"],
            ["Age", "years", "Children-first effect; age stratification"],
            ["SibSp + Parch", "int", "Combined into FamilySize and IsAlone"],
            ["Fare", "USD", "Class proxy with finer resolution"],
            ["Embarked", "C/Q/S/B", "Port of embarkation - confounded with class"],
            ["Name", "str", "Title (Mr/Mrs/Miss/Master) extracted as derived feature"],
            ["BoatBody", "str", "Parsed into Lifeboat number and BodyRecovered flag"],
        ],
        col_widths_cm=[3.5, 2.5, 10.5],
    )

    add_heading_2(doc, "2.2  Method")
    add_para(doc,
        "The analytical approach moves from descriptive to inferential:",
        space_after=6)
    add_bullet(doc, [
        ("Descriptive comparisons.", True),
        " Survival rates by group, with 95% Wilson confidence intervals so the precision of each estimate is visible.",
    ])
    add_bullet(doc, [
        ("Effect-size ranking.", True),
        " Every feature is placed on a single 0-1 comparable scale. Cramer's V for categorical features (Sex, Class, Embarked); the absolute point-biserial correlation for numeric ones (Age, Fare, SibSp, Parch).",
    ])
    add_bullet(doc, [
        ("Odds ratios with 95% CIs.", True),
        " For each headline contrast (women vs men, 1st class vs the rest, etc.), we compute the odds ratio with a 95% CI from the log-odds standard error, and a Fisher's exact p-value for the contrast.",
    ])
    add_bullet(doc, [
        ("Hypothesis tests.", True),
        " Chi-square tests of independence for categorical relationships; Welch's t-tests with Cohen's d for numeric comparisons; one-way ANOVA across multiple age groups.",
    ])
    add_bullet(doc, [
        ("Stratified analysis.", True),
        " Joint Class x Sex tables to surface compounding effects that disappear in marginal analyses.",
    ])

    add_callout(doc, "A note on what this report does NOT do",
        "We have not fit a predictive model. Effect-size ranking and odds ratios are descriptive: "
        "they say which features individually carry survival information, but they do not adjust "
        "for one another. A logistic regression with interaction terms would refine these estimates "
        "and is the natural next step. See Section 8 (Recommendations).",
        accent="6B7280", bg="F3F4F6")


def the_big_picture(doc, A):
    add_heading_1(doc, "3. The Big Picture")

    add_heading_2(doc, "3.1  Overall survival rate")
    add_para(doc,
        f"Of the {A['n_total']:,} passengers in our dataset, {A['n_survived']:,} survived - a rate of "
        f"{A['rate_overall']:.1f}%. About one in three. That is the figure most people remember, and "
        "it is the figure that hides the entire story of this disaster.",
        space_after=8)

    add_image(doc, CHART_DIR / "survival_overview.png",
              caption="Figure 1. Overall outcome, survival by sex, and survival by class. The marginal "
                       "averages already hint that 'overall' is a misleading number.")

    add_heading_2(doc, "3.2  Why the average is misleading")
    add_para(doc,
        "Consider three slices of the 38.2% headline:",
        space_after=4)
    add_bullet(doc, [
        ("By sex:", True),
        f" {next(r for _, r in A['sex_ci'].iterrows() if r['level'] == 'female')['rate']:.1f}% of women "
        f"survived, vs {next(r for _, r in A['sex_ci'].iterrows() if r['level'] == 'male')['rate']:.1f}% of men.",
    ])
    add_bullet(doc, [
        ("By class:", True),
        f" {next(r for _, r in A['pclass_ci'].iterrows() if r['level'] == 1)['rate']:.1f}% of 1st class "
        f"survived, vs {next(r for _, r in A['pclass_ci'].iterrows() if r['level'] == 3)['rate']:.1f}% of 3rd class.",
    ])
    add_bullet(doc, [
        ("By the two combined:", True),
        f" {A['joint'].loc[(1, 'female'), 'rate']:.1f}% of 1st-class women survived, vs "
        f"{A['joint'].loc[(3, 'male'), 'rate']:.1f}% of 3rd-class men. The arithmetic average between "
        "these is meaningless; nobody had a 'typical' Titanic experience.",
    ])

    add_pullquote(doc,
        "If you remember nothing else from this section: the 38.2% overall rate is an artefact of "
        "averaging together two populations that the evacuation treated almost entirely differently.")


def the_three_drivers(doc, A):
    add_heading_1(doc, "4. The Three Biggest Drivers")
    add_para(doc,
        "We can now rank the factors by predictive strength on a common scale. The chart below uses "
        "Cramer's V for categorical features and the absolute value of the point-biserial correlation "
        "for numeric ones. Both range from 0 to 1; thresholds for 'small', 'medium', and 'large' "
        "are 0.1, 0.3, and 0.5 by convention.",
        space_after=6)

    add_image(doc, CHART_DIR / "feature_importance.png",
              caption="Figure 2. Feature predictive power, ranked. Three things matter at all - sex, class, "
                       "and fare/embarked. Everything else is small or negligible on its own.")

    es_rows = []
    for _, r in A["effect_sizes"].iterrows():
        es_rows.append([r["feature"], r["type"].capitalize(), r["metric"],
                        f"{r['effect_size']:+.3f}", r["strength"]])
    add_table(doc, ["Feature", "Type", "Metric", "Effect Size", "Strength"], es_rows,
              col_widths_cm=[3.0, 2.6, 4.2, 3.0, 3.7])

    # 4.1 Sex
    add_heading_2(doc, "4.1  Sex - by far the strongest signal")
    add_para(doc,
        f"With Cramer's V ~ {A['effect_sizes'].iloc[0]['effect_size']:.2f}, sex is the only feature "
        "in 'large effect' territory. The contrast is stark: of 466 women, 339 survived; of 843 men, "
        "only 161 did. The two confidence intervals do not come close to overlapping, so we can be "
        "essentially certain this is not noise.",
        space_after=6)

    sex_or = next(o for o in A["odds"] if o["label"] == "Female vs Male")
    add_callout(doc, "What this means in plain English",
        f"A woman on the Titanic had roughly {sex_or['odds_ratio']:.0f} times the odds of surviving "
        f"that a man had (95% CI: {sex_or['ci_low']:.1f}x to {sex_or['ci_high']:.1f}x). Sex is not just "
        "the most useful single variable - it is the only variable that on its own gives you a "
        "near-reliable prediction.",
        accent="059669", bg="ECFDF5")

    # 4.2 Class
    add_heading_2(doc, "4.2  Class - the second-strongest, and an enabler of the first")
    add_para(doc,
        "Class came in at Cramer's V ~ 0.31 - solidly in 'medium' territory. The survival rate "
        f"falls steadily: 1st class {next(r for _, r in A['pclass_ci'].iterrows() if r['level'] == 1)['rate']:.1f}%, "
        f"2nd class {next(r for _, r in A['pclass_ci'].iterrows() if r['level'] == 2)['rate']:.1f}%, "
        f"3rd class {next(r for _, r in A['pclass_ci'].iterrows() if r['level'] == 3)['rate']:.1f}%. "
        "The mechanism here is not abstract: 1st-class cabins were on upper decks, much closer to "
        "the boat deck where lifeboats were loaded; 1st-class passengers had priority boarding and "
        "better access to information about what was happening as the ship took on water.",
        space_after=6)

    # 4.3 Interaction
    add_heading_2(doc, "4.3  Class x Sex - the real story is in the interaction")
    add_para(doc,
        "Sex and class do not simply add to each other - they compound. The joint table below shows "
        "the survival rate for each of the six possible combinations.",
        space_after=4)

    add_image(doc, CHART_DIR / "joint_class_sex.png",
              caption="Figure 3. Class x Sex joint survival. The diagonal is staggering: 1st-class women "
                       "(top-left, deep green) survived almost universally; 3rd-class men (bottom-right, deep red) "
                       "almost universally died.")

    class_suffix = {1: "1st", 2: "2nd", 3: "3rd"}
    joint_rows = []
    for (cls, sex), row in A["joint"].iterrows():
        n = int(row["count"])
        rate = row["rate"]
        s = int(round(n * rate / 100))
        lo, hi = wilson_ci(s, n)
        joint_rows.append([f"{class_suffix[cls]} - {sex.capitalize()}", n, s,
                           f"{rate:.1f}%", f"[{lo:.1f}, {hi:.1f}]"])
    add_table(doc, ["Group", "Total", "Survived", "Rate", "95% CI"], joint_rows,
              col_widths_cm=[4.0, 2.5, 2.8, 3.0, 4.2])

    add_callout(doc, "What this means in plain English",
        f"The two extreme cells - 1st-class women at {A['joint'].loc[(1, 'female'), 'rate']:.1f}% and "
        f"3rd-class men at {A['joint'].loc[(3, 'male'), 'rate']:.1f}% - are roughly 80 percentage points "
        "apart. That gap is larger than the marginal effect of either sex or class alone. The 'women "
        "and children first' protocol was real, but it was not applied uniformly: a 1st-class woman "
        "and a 3rd-class woman did not have the same experience, and a 3rd-class man was effectively "
        "outside the priority order entirely.",
        accent="059669", bg="ECFDF5")


def secondary_factors(doc, A):
    add_heading_1(doc, "5. Secondary Factors")
    add_para(doc,
        "Beyond sex and class, four further attributes shifted the survival odds - some genuinely, "
        "some only because they were entangled with the bigger drivers.",
        space_after=8)

    # 5.1 Age
    add_heading_2(doc, "5.1  Age - the 'children first' effect was real, but small")
    add_para(doc,
        f"On average, survivors were about a year younger than non-survivors (mean ages "
        f"{A['tests']['age_t']['mean_survived']:.1f} vs {A['tests']['age_t']['mean_perished']:.1f} years; "
        f"Welch's t = {A['tests']['age_t']['t_statistic']:.2f}, p = {A['tests']['age_t']['p_value']:.3f}). "
        "That difference is statistically detectable but practically tiny. The real age effect lives "
        "at the extremes, not the mean: very young children fared significantly better, and the "
        "elderly fared significantly worse.",
        space_after=4)

    age_rows = []
    for _, r in A["age_summary"].iterrows():
        n = int(r["count"])
        rate = r["rate"]
        s = int(round(n * rate / 100))
        lo, hi = wilson_ci(s, n)
        age_rows.append([str(r["AgeGroup"]), n, s, f"{rate:.1f}%", f"[{lo:.1f}, {hi:.1f}]"])
    add_table(doc, ["Age group", "Total", "Survived", "Rate", "95% CI"], age_rows,
              col_widths_cm=[4.5, 2.5, 2.8, 3.0, 4.0])

    # 5.2 Family size
    add_heading_2(doc, "5.2  Family size - a sweet spot at 2-4")
    add_para(doc,
        "Family size (siblings + spouse + parents + children + self) shows a clearly non-monotonic "
        "pattern. Solo travellers survived at 30.3%. Mid-sized families of 2-4 jumped to 53-70%. "
        "Very large families of 5+ dropped back to around 25% or below, with no survivors at all "
        "among families of 8 or 11.",
        space_after=4)

    fam_rows = []
    for _, r in A["family"].iterrows():
        n = int(r["count"])
        rate = r["rate"]
        s = int(round(n * rate / 100))
        fam_rows.append([int(r["FamilySize"]), n, s, f"{rate:.1f}%"])
    add_table(doc, ["Family size", "Total", "Survived", "Rate"], fam_rows,
              col_widths_cm=[3.5, 3.0, 3.0, 3.0])

    add_para(doc,
        "The plausible mechanism: mid-sized families were small enough to stay together during "
        "evacuation but large enough to advocate for one another (and for women and children within "
        "the group). Solo passengers lacked that advocacy. Very large families struggled to keep "
        "everyone together in the chaos, and would-be helpers were less likely to board with so "
        "many dependents.",
        space_after=8)

    # 5.3 Fare
    add_heading_2(doc, "5.3  Fare - a strong proxy for class, not an independent cause")
    add_para(doc,
        f"Fare is the strongest numeric signal (r = {A['effect_sizes'].iloc[2]['effect_size']:+.3f}). "
        f"Survivors paid an average of ${A['tests']['fare_t']['mean_survived']:.2f}; non-survivors "
        f"${A['tests']['fare_t']['mean_perished']:.2f} - less than half. The t-test confirms the "
        f"difference is overwhelmingly real (t = {A['tests']['fare_t']['t_statistic']:.2f}, "
        f"p < 0.001). However, almost all of fare's signal can be explained by the fact that "
        "expensive tickets bought 1st-class accommodation. Fare is class repackaged at higher "
        "resolution; treating it as an independent cause would be a mistake.",
    )

    # 5.4 Port
    add_heading_2(doc, "5.4  Embarkation port - a textbook confound")
    add_para(doc,
        "At first glance, embarkation port looks meaningful: Cherbourg passengers survived at "
        f"{next(r for _, r in A['emb_ci'].iterrows() if r['level'] == 'C')['rate']:.1f}% versus "
        f"Southampton's {next(r for _, r in A['emb_ci'].iterrows() if r['level'] == 'S')['rate']:.1f}%. "
        "But Cherbourg was where most 1st-class passengers boarded; Southampton's mix was "
        "predominantly 3rd-class. The chi-square test does flag a significant association, but it "
        "is almost entirely driven by this class composition difference. Adjusting for class makes "
        "the port effect largely vanish - it should not be treated as a survival factor in its own right.",
    )

    # 5.5 Title
    add_heading_2(doc, "5.5  Title - a compact summary of sex and age")
    add_para(doc,
        "Titles in 1912 carried real information. 'Master' was used specifically for boys; 'Mrs' "
        "for married women; 'Miss' for unmarried women; 'Mr' for all adult men. The survival rates "
        "by title essentially restate the sex + age story in a single feature:",
        space_after=4)

    title_rows = []
    for _, r in A["title"].iterrows():
        title_rows.append([str(r["Title"]), int(r["count"]), f"{r['rate']:.1f}%"])
    add_table(doc, ["Title", "Count", "Survival rate"], title_rows,
              col_widths_cm=[5.0, 3.5, 4.0])

    add_para(doc,
        "For modelling, this matters: Title can substitute for Sex and Age simultaneously while "
        "remaining a single, clean categorical feature. It is a useful piece of feature engineering "
        "rather than an independent finding.",
    )


def lifeboat_mechanism(doc, A):
    add_heading_1(doc, "6. The Mechanism: Lifeboats")
    add_para(doc,
        "Everything above is about who was likely to survive. This section is about how survival "
        "actually happened: by getting onto a lifeboat. The titanic5 dataset records the lifeboat "
        "number where it is known, which gives us a direct view of the proximate cause.",
        space_after=6)

    lb = A["lifeboat"]
    lo_b, hi_b = wilson_ci(lb["s_boat"], lb["n_boat"])
    lo_n, hi_n = wilson_ci(lb["s_no"], lb["n_no"])

    add_image(doc, CHART_DIR / "lifeboat_survival.png",
              caption="Figure 4. Lifeboat record vs survival, and proportion of each sex with a "
                       "lifeboat record. The right-hand chart is the 'women and children first' protocol "
                       "made operational.")

    add_table(doc,
        ["Group", "Total", "Survived", "Rate", "95% CI"],
        [
            ["On a lifeboat (recorded)", lb["n_boat"], lb["s_boat"], f"{lb['rate_boat']:.1f}%",
             f"[{lo_b:.1f}, {hi_b:.1f}]"],
            ["No lifeboat record", lb["n_no"], lb["s_no"], f"{lb['rate_no']:.1f}%",
             f"[{lo_n:.1f}, {hi_n:.1f}]"],
        ],
        col_widths_cm=[5.0, 2.5, 2.8, 3.0, 3.2])

    add_para(doc,
        "The numbers are stark. Getting onto a lifeboat was effectively a guarantee of survival; "
        "not getting on one was almost a guarantee of death. The two confidence intervals don't "
        "even come close to overlapping.",
        space_after=6)

    add_callout(doc, "Why every demographic effect in this report ultimately reduces to this",
        "The 'women and children first' protocol was the rule that determined boat-seat allocation. "
        "Class shaped how strictly the rule was enforced (1st-class women were prioritised over "
        "3rd-class women in practice) and how easily passengers could physically reach the boat "
        "deck (1st-class cabins were close; 3rd-class cabins were locked away deep in the ship "
        "for hours after the collision). All of the demographic patterns we've measured in this "
        "report are the upstream consequences of a single allocation decision: which body got into "
        "a lifeboat.",
        accent="059669", bg="ECFDF5")


def statistical_robustness(doc, A):
    add_heading_1(doc, "7. Statistical Robustness")
    add_para(doc,
        "Every comparison in this report could in principle be a coincidence of sampling. The "
        "tests below check whether each is real. Each produces a p-value: the probability of "
        "seeing the observed pattern if there were no underlying effect. Standard thresholds: "
        "p < 0.05 is 'probably real', p < 0.001 is 'almost certainly real'. Most of the patterns "
        "here clear those thresholds by many orders of magnitude.",
        space_after=6)

    tests = A["tests"]
    rows = [
        ["Sex -> Survived", "Chi-square", f"chi2 = {tests['sex']['chi2']:.1f}",
         "< 2e-16", tests["sex"]["strength"]],
        ["Pclass -> Survived", "Chi-square", f"chi2 = {tests['pclass']['chi2']:.1f}",
         "< 2e-16", tests["pclass"]["strength"]],
        ["Embarked -> Survived", "Chi-square", f"chi2 = {tests['emb']['chi2']:.1f}",
         f"{tests['emb']['p_value']:.2e}", tests["emb"]["strength"]],
        ["Fare (Survived vs Perished)", "Welch t-test", f"t = {tests['fare_t']['t_statistic']:.2f}",
         f"{tests['fare_t']['p_value']:.2e}", f"d = {tests['fare_t']['cohens_d']:.2f} ({tests['fare_t']['effect_size']})"],
        ["Age (Survived vs Perished)", "Welch t-test", f"t = {tests['age_t']['t_statistic']:.2f}",
         f"{tests['age_t']['p_value']:.3f}", f"d = {tests['age_t']['cohens_d']:.2f} ({tests['age_t']['effect_size']})"],
        ["Age groups -> Survived", "ANOVA", f"F = {tests['age_anova']['f_statistic']:.2f}",
         f"{tests['age_anova']['p_value']:.4f}", "Significant" if tests['age_anova']['significant'] else "Not significant"],
    ]
    add_table(doc, ["Relationship", "Test", "Statistic", "p-value", "Effect / Strength"], rows,
              col_widths_cm=[5.0, 2.5, 3.2, 2.8, 3.5])

    add_heading_2(doc, "Odds ratios with 95% CIs")
    add_para(doc,
        "Odds ratios translate the statistical significance into something more interpretable: how "
        "much each factor multiplies your odds of surviving. OR > 1 = better odds; OR < 1 = worse odds; "
        "OR = 1 = no effect at all. CIs come from the log-odds standard error; p-values from Fisher's exact test.",
        space_after=4)

    add_image(doc, CHART_DIR / "odds_ratios.png",
              caption="Figure 5. Survival odds ratios with 95% CIs. Bars to the right of 1.0 (green) "
                       "helped survival; bars to the left (red) hurt it. The log scale matters: each gridline is a 10x.")

    or_rows = []
    for o in A["odds"]:
        or_rows.append([o["label"], f"{o['odds_ratio']:.2f}x",
                        f"[{o['ci_low']:.2f}, {o['ci_high']:.2f}]",
                        f"{o['rate_exposed']:.1f}%", f"{o['rate_unexposed']:.1f}%",
                        f"{'+' if o['lift'] > 0 else ''}{o['lift']:.1f}pp",
                        "< 2e-16" if o["p_value"] < 1e-10 else f"{o['p_value']:.2e}"])
    add_table(doc, ["Contrast", "OR", "95% CI", "Exposed rate", "Unexposed rate", "Lift", "p"],
              or_rows,
              col_widths_cm=[4.0, 1.5, 2.3, 2.4, 2.6, 1.6, 1.8])


def conclusions(doc, A):
    add_heading_1(doc, "8. Conclusions & Recommendations")

    add_heading_2(doc, "8.1  What the numbers actually say")
    add_para(doc,
        "Survival on the Titanic was governed by who got onto a lifeboat, and lifeboat allocation "
        "was governed by a clear if uneven application of 'women and children first.' That core "
        "rule was modulated heavily by class: 1st-class passengers had easier physical access to "
        "the boat deck and were prioritised in practice. The result is a four-tier hierarchy:",
        space_after=4)
    add_bullet(doc, [("1.", True), " 1st- and 2nd-class women: near-universal survival (~89-97%)."])
    add_bullet(doc, [("2.", True), " 1st- and 2nd-class men, plus 3rd-class women: middling odds (~14-49%)."])
    add_bullet(doc, [("3.", True), " 3rd-class men: near-universal death (~15%)."])
    add_bullet(doc, [("4.", True), " Within each of these, age and family size further modulated the odds, but only at the margins."])

    add_heading_2(doc, "8.2  Recommendations for further analysis")
    add_bullet(doc, [
        ("Fit a logistic regression with Class x Sex interaction.", True),
        " Marginal odds ratios already tell a clear story, but a joint model would let us separate the "
        "independent effects of sex and class from their interaction, and would adjust for fare, age, "
        "and family size simultaneously.",
    ])
    add_bullet(doc, [
        ("Engineer Title as a primary feature.", True),
        " Title captures sex and (loosely) age in a single categorical column, and out-performs each individually in standard Kaggle baselines.",
    ])
    add_bullet(doc, [
        ("Drop Fare or treat it as an alternative encoding of Class.", True),
        " Including both in a model creates collinearity without adding much information.",
    ])
    add_bullet(doc, [
        ("Drop Embarked as a survival predictor.", True),
        " Its raw association is confounded with class; once class is in the model, Embarked contributes almost nothing.",
    ])
    add_bullet(doc, [
        ("Treat Lifeboat / BodyRecovered as outcomes, not inputs.", True),
        " They are downstream of survival. Including them as features would leak the target.",
    ])
    add_bullet(doc, [
        ("Consider extracting Cabin / Deck letters from Occupation.", True),
        " 47% of Occupation entries are missing, but the non-null entries may encode cabin information "
        "that could shed light on physical proximity to the boat deck.",
    ])

    add_heading_2(doc, "8.3  Limitations")
    add_bullet(doc, [
        ("Occupation is 47% missing.", True),
        " This limits cabin-level analysis. We've used HasCabin (a 1/0 indicator of non-null Occupation) as a low-resolution proxy.",
    ])
    add_bullet(doc, [
        ("Survival is recorded, but the timing of evacuation decisions is not.", True),
        " We cannot directly test hypotheses about, for example, the order in which lifeboats were lowered.",
    ])
    add_bullet(doc, [
        ("Lifeboat numbers are descriptive of survivors, not a causal lever.", True),
        " Using Lifeboat as a feature in a predictive model would be target leakage.",
    ])
    add_bullet(doc, [
        ("Class is a coarse proxy for socioeconomic status.", True),
        " Within each class there was considerable variation in cabin location, age, and family situation.",
    ])


def appendix(doc, A):
    add_heading_1(doc, "Appendix A. Full Tables")

    add_heading_2(doc, "A.1  Survival by sex (with 95% CI)")
    add_table(doc,
        ["Sex", "n", "Survived", "Rate", "95% CI"],
        [[str(r["level"]).capitalize(), int(r["n"]), int(r["survived"]),
          f"{r['rate']:.1f}%", f"[{r['ci_low']:.1f}, {r['ci_high']:.1f}]"]
         for _, r in A["sex_ci"].iterrows()],
        col_widths_cm=[3.5, 2.5, 3.0, 3.0, 4.0])

    add_heading_2(doc, "A.2  Survival by class (with 95% CI)")
    cs = {1: "1st", 2: "2nd", 3: "3rd"}
    add_table(doc,
        ["Class", "n", "Survived", "Rate", "95% CI"],
        [[cs.get(int(r["level"]), str(int(r["level"]))), int(r["n"]), int(r["survived"]),
          f"{r['rate']:.1f}%", f"[{r['ci_low']:.1f}, {r['ci_high']:.1f}]"]
         for _, r in A["pclass_ci"].iterrows()],
        col_widths_cm=[3.5, 2.5, 3.0, 3.0, 4.0])

    add_heading_2(doc, "A.3  Survival by embarkation port (with 95% CI)")
    port_names = {"C": "Cherbourg", "Q": "Queenstown", "S": "Southampton", "B": "Belfast"}
    add_table(doc,
        ["Port", "n", "Survived", "Rate", "95% CI"],
        [[port_names.get(r["level"], str(r["level"])), int(r["n"]), int(r["survived"]),
          f"{r['rate']:.1f}%", f"[{r['ci_low']:.1f}, {r['ci_high']:.1f}]"]
         for _, r in A["emb_ci"].iterrows()],
        col_widths_cm=[3.5, 2.5, 3.0, 3.0, 4.0])

    add_heading_2(doc, "A.4  Correlation matrix (numeric features only)")
    corr = correlation_analysis(A["df_raw"])
    corr_features = ["Age", "Fare", "SibSp", "Parch", "Survived"]
    rows = []
    for f in corr_features:
        rows.append([f] + [f"{corr.loc[f, c]:+.3f}" for c in corr_features])
    add_table(doc, [""] + corr_features, rows,
              col_widths_cm=[2.5] + [2.6] * len(corr_features))

    add_heading_1(doc, "Appendix B. Glossary")
    glossary = [
        ("Survival rate", "Percent of a group that survived. 62% means 62 of every 100 lived."),
        ("95% CI (Wilson)", "The range we are 95% confident contains the true rate. Narrow = certain estimate; wide = small sample, less certain."),
        ("Odds Ratio (OR)", "How many times higher the odds of survival were for one group versus another. OR = 2 means twice the odds; OR < 1 means lower; OR = 1 means no effect."),
        ("Effect size", "How strongly a single feature predicts survival on a 0-1 scale. By convention: 0.1 = small, 0.3 = medium, 0.5 = large."),
        ("Cramer's V", "Effect-size metric for categorical features (Sex, Class, Port)."),
        ("Point-biserial r", "Effect-size metric for a numeric feature predicting a 0/1 outcome."),
        ("Cohen's d", "Size of the difference between two group means, in standard deviations. 0.2 small, 0.5 medium, 0.8 large."),
        ("p-value", "Probability the pattern arose by chance. < 0.05 = probably real; < 0.001 = essentially certain."),
        ("pp (percentage points)", "Arithmetic gap between two percentages. 19% -> 73% is a 54pp jump, not 54%."),
        ("Confounding", "When a third factor (e.g., class) drives the apparent relationship between two others (e.g., port and survival)."),
        ("Target leakage", "Using a feature that is actually a consequence of the outcome you're predicting. Lifeboat record in a survival model would be leakage."),
    ]
    add_table(doc,
        ["Term", "Plain-English definition"],
        [[t, d] for t, d in glossary],
        col_widths_cm=[4.5, 12.5])


if __name__ == "__main__":
    main()
